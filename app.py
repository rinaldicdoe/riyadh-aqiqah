import streamlit as st
import pandas as pd
import numpy as np
import io
import openpyxl 
import datetime
import os
import re
import shutil
import zipfile
from zipfile import ZipFile
import tempfile
# Import library untuk Word
from docx import Document
from docx.shared import Cm, Pt, RGBColor

# Initialize session state
if 'missing_items' not in st.session_state:
    st.session_state.missing_items = None

# ============================================================
# HELPER FUNCTIONS: EXCEL READING WITH ROBUST HANDLING
# ============================================================
def repair_xlsx_file(file_path):
    """
    Memperbaiki file XLSX dengan XML yang invalid dengan cara:
    1. Ekstrak file XLSX (yang merupakan ZIP)
    2. Perbaiki XML yang invalid di SEMUA file
    3. Kembalikan file yang diperbaiki
    """
    try:
        # Coba import lxml jika tersedia
        try:
            from lxml import etree
            use_lxml = True
        except ImportError:
            use_lxml = False
        
        # Buat direktori sementara untuk ekstraksi
        extract_dir = tempfile.mkdtemp()
        repaired_path = file_path + ".repaired.xlsx"
        
        try:
            # Ekstrak XLSX
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(extract_dir)
            
            # Cari SEMUA file XML di dalam XLSX
            xl_dir = os.path.join(extract_dir, 'xl')
            if os.path.exists(xl_dir):
                for root, dirs, files in os.walk(xl_dir):
                    for file in files:
                        if file.endswith('.xml') or file.endswith('.rels'):
                            xml_file = os.path.join(root, file)
                            try:
                                with open(xml_file, 'rb') as f:
                                    content = f.read()
                                
                                # Hapus karakter invalid
                                content = content.replace(b'\x00', b'')
                                
                                # Hapus karakter kontrol yang tidak diizinkan XML (kecuali tab, newline, carriage return)
                                content = re.sub(b'[\x00-\x08\x0b\x0c\x0e-\x1f]', b'', content)
                                
                                if use_lxml:
                                    try:
                                        from lxml import etree
                                        parser = etree.XMLParser(recover=True, remove_blank_text=False)
                                        tree = etree.fromstring(content, parser)
                                        with open(xml_file, 'wb') as f:
                                            f.write(etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True))
                                    except:
                                        # Tulis kembali content yang sudah dibersihkan
                                        with open(xml_file, 'wb') as f:
                                            f.write(content)
                                else:
                                    # Jika lxml tidak ada, tulis kembali content yang sudah dibersihkan
                                    with open(xml_file, 'wb') as f:
                                        f.write(content)
                            except Exception as e:
                                pass
            
            # Rekompres XLSX
            with zipfile.ZipFile(repaired_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, dirs, files in os.walk(extract_dir):
                    for file in files:
                        file_path_full = os.path.join(root, file)
                        arcname = os.path.relpath(file_path_full, extract_dir)
                        zipf.write(file_path_full, arcname)
            
            return repaired_path
        finally:
            # Bersihkan direktori sementara
            shutil.rmtree(extract_dir, ignore_errors=True)
    except Exception as e:
        return None

def read_excel_robust(uploaded_file, header=None):
    """
    Membaca file Excel dengan penanganan untuk file yang belum di-enable editing.
    Jika pembacaan normal gagal, coba beberapa metode alternatif.
    Selalu membaca sheet pertama (index 0).
    """
    filename = uploaded_file.name
    last_error = None
    tmp_files_to_clean = []
    
    try:
        # Metode 1: Pembacaan normal dengan data_only=False
        try:
            if filename.endswith('.xlsx'):
                uploaded_file.seek(0)
                df = pd.read_excel(uploaded_file, sheet_name=0, header=header, engine='openpyxl', engine_kwargs={'data_only': False})
                return df
        except Exception as e:
            last_error = str(e)
        
        # Metode 2: Simpan ke file sementara dan coba dengan load_workbook yang lebih permisif
        try:
            uploaded_file.seek(0)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
                tmp_file.write(uploaded_file.getbuffer())
                tmp_path = tmp_file.name
            tmp_files_to_clean.append(tmp_path)
            
            # Coba membuka workbook dengan keep_vba=False untuk menghindari XML issues
            wb = openpyxl.load_workbook(tmp_path, keep_vba=False, data_only=False)
            
            # Ambil sheet pertama
            ws = wb.worksheets[0]
            
            # Baca data dari worksheet
            data = []
            for row in ws.iter_rows(values_only=True):
                data.append(row)
            
            # Tentukan header
            if header is None:
                df = pd.DataFrame(data)
            elif isinstance(header, int):
                df = pd.DataFrame(data[header+1:])
                df.columns = data[header]
            else:
                df = pd.DataFrame(data)
            
            return df
        except Exception as e:
            last_error = str(e)
        
        # Metode 3: Perbaiki file XLSX terlebih dahulu, kemudian baca
        try:
            uploaded_file.seek(0)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
                tmp_file.write(uploaded_file.getbuffer())
                tmp_path = tmp_file.name
            tmp_files_to_clean.append(tmp_path)
            
            # Perbaiki file XLSX
            repaired_path = repair_xlsx_file(tmp_path)
            if repaired_path:
                tmp_files_to_clean.append(repaired_path)
                
                try:
                    # Coba baca file yang sudah diperbaiki dengan openpyxl
                    wb = openpyxl.load_workbook(repaired_path, keep_vba=False, data_only=False)
                    ws = wb.worksheets[0]
                    data = []
                    for row in ws.iter_rows(values_only=True):
                        data.append(row)
                    
                    if header is None:
                        df = pd.DataFrame(data)
                    elif isinstance(header, int):
                        df = pd.DataFrame(data[header+1:])
                        df.columns = data[header]
                    else:
                        df = pd.DataFrame(data)
                    
                    return df
                except Exception as e:
                    last_error = str(e)
                
                # Jika openpyxl masih gagal, coba pandas dengan repaired file
                try:
                    df = pd.read_excel(repaired_path, sheet_name=0, header=header, engine='openpyxl')
                    return df
                except Exception as e:
                    last_error = str(e)
        except Exception as e:
            last_error = str(e)
        
        # Metode 4: Coba dengan engine default pandas pada file original
        try:
            uploaded_file.seek(0)
            df = pd.read_excel(uploaded_file, sheet_name=0, header=header)
            return df
        except Exception as e:
            last_error = str(e)
        
        # Metode 5: Ekstrak styles.xml yang corrupt dan baca tanpa styling
        try:
            uploaded_file.seek(0)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
                tmp_file.write(uploaded_file.getbuffer())
                tmp_path = tmp_file.name
            tmp_files_to_clean.append(tmp_path)
            
            # Ekstrak dan hapus styles.xml yang corrupt
            extract_dir = tempfile.mkdtemp()
            
            with zipfile.ZipFile(tmp_path, 'r') as zip_ref:
                zip_ref.extractall(extract_dir)
            
            styles_path = os.path.join(extract_dir, 'xl', 'styles.xml')
            if os.path.exists(styles_path):
                try:
                    os.remove(styles_path)
                except:
                    pass
            
            # Rekompres tanpa styles.xml
            new_tmp_path = tmp_path + ".no_styles.xlsx"
            with zipfile.ZipFile(new_tmp_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, dirs, files in os.walk(extract_dir):
                    for file in files:
                        file_path_full = os.path.join(root, file)
                        arcname = os.path.relpath(file_path_full, extract_dir)
                        zipf.write(file_path_full, arcname)
            
            tmp_files_to_clean.append(new_tmp_path)
            shutil.rmtree(extract_dir, ignore_errors=True)
            
            # Coba baca file tanpa styles
            try:
                wb = openpyxl.load_workbook(new_tmp_path, keep_vba=False, data_only=False)
                ws = wb.worksheets[0]
                data = []
                for row in ws.iter_rows(values_only=True):
                    data.append(row)
                
                if header is None:
                    df = pd.DataFrame(data)
                elif isinstance(header, int):
                    df = pd.DataFrame(data[header+1:])
                    df.columns = data[header]
                else:
                    df = pd.DataFrame(data)
                
                return df
            except Exception as e:
                last_error = str(e)
                try:
                    df = pd.read_excel(new_tmp_path, sheet_name=0, header=header)
                    return df
                except Exception as e:
                    last_error = str(e)
        except Exception as e:
            last_error = str(e)
        
        # Jika semua metode gagal, lempar error
        raise Exception(f"Tidak dapat membaca file Excel. File mungkin corrupt atau formatnya tidak valid. Silakan coba: 1) Buka file di Excel dan Save As, 2) Atau unggah sebagai CSV. Error: {last_error}")
    
    finally:
        # Bersihkan file sementara
        for tmp_file in tmp_files_to_clean:
            try:
                if os.path.exists(tmp_file):
                    os.unlink(tmp_file)
            except:
                pass

# ============================================================
# HELPER FUNCTIONS: EXCEL FORMATTING
# ============================================================
def format_rekap_pemotongan_excel(writer, df):
    """
    Format Excel untuk Rekap Pemotongan dengan:
    - Font size 22 untuk semua cell
    - Table format
    - Column width yang sesuai
    - Warna untuk kata-kata tertentu
    """
    workbook = writer.book
    worksheet = writer.sheets['Hasil Rekap Pemotongan']
    
    # Define base formats
    header_format = workbook.add_format({
        'font_size': 22,
        'bold': True,
        'valign': 'vcenter',
        'align': 'center',
        'text_wrap': True,
        'border': 1,
        'bg_color': '#D3D3D3'
    })
    
    cell_format = workbook.add_format({
        'font_size': 22,
        'valign': 'vcenter',
        'align': 'left',
        'text_wrap': True,
        'border': 1
    })
    
    cell_center_format = workbook.add_format({
        'font_size': 22,
        'valign': 'vcenter',
        'align': 'center',
        'text_wrap': True,
        'border': 1
    })
    
    # Date format untuk Tanggal Kirim dan Tanggal Potong
    # Date format sebagai TEXT dengan DD-MM-YYYY (gunakan custom format Excel)
    date_format = workbook.add_format({
        'font_size': 22,
        'valign': 'vcenter',
        'align': 'center',
        'text_wrap': True,
        'border': 1,
        'num_format': 'dd/mm/yyyy'  # Explicitly DD/MM/YYYY
    })
    
    # Format warna untuk Tanggal Kirim = Tanggal Potong (dark yellow background)
    date_same_format = workbook.add_format({
        'font_size': 22,
        'font_color': '#000000',  # Black text
        'bg_color': "#B74706",  # Dark yellow background
        'bold': True,
        'valign': 'vcenter',
        'align': 'center',
        'text_wrap': True,
        'border': 1,
        'num_format': 'dd/mm/yyyy'  # Explicitly DD/MM/YYYY
    })
    
    # Format warna background untuk kolom M (Paket)
    paket_jantan_format = workbook.add_format({
        'font_size': 22,
        'font_color': '#000000',  # Black text
        'bg_color': '#00B050',  # Green background
        'bold': True,
        'valign': 'vcenter',
        'align': 'left',
        'text_wrap': True,
        'border': 1
    })
    
    paket_kebuli_format = workbook.add_format({
        'font_size': 22,
        'font_color': '#000000',  # Black text
        'bg_color': '#FF8C00',  # Orange background
        'bold': True,
        'valign': 'vcenter',
        'align': 'left',
        'text_wrap': True,
        'border': 1
    })
    
    # Format warna background untuk kolom P (Pemotongan Disaksikan)
    pemotongan_live_format = workbook.add_format({
        'font_size': 22,
        'font_color': '#000000',  # Black text
        'bg_color': '#0070C0',  # Blue background
        'bold': True,
        'valign': 'vcenter',
        'align': 'left',
        'text_wrap': True,
        'border': 1
    })
    
    pemotongan_disaksikan_format = workbook.add_format({
        'font_size': 22,
        'font_color': "#000000",  # Black text
        'bg_color': "#CA6F00",  # Yellow background
        'bold': True,
        'valign': 'vcenter',
        'align': 'left',
        'text_wrap': True,
        'border': 1
    })
    
    # Format warna background untuk kolom Q (Catatan Khusus)
    catatan_yellow_format = workbook.add_format({
        'font_size': 22,
        'font_color': "#000000",  # Black text
        'bg_color': "#B79602",  # Yellow background
        'bold': True,
        'valign': 'vcenter',
        'align': 'left',
        'text_wrap': True,
        'border': 1
    })
    
    # Format untuk baris dengan Status "Belum Dikonfirmasi" (red background with BLACK text)
    belum_dikonfirmasi_format = workbook.add_format({
        'font_size': 22,
        'font_color': '#000000',  # Black text (changed from white)
        'bg_color': '#FF0000',  # Red background
        'bold': True,
        'valign': 'vcenter',
        'align': 'left',
        'text_wrap': True,
        'border': 1
    })
    
    belum_dikonfirmasi_center_format = workbook.add_format({
        'font_size': 22,
        'font_color': '#000000',  # Black text (changed from white)
        'bg_color': '#FF0000',  # Red background
        'bold': True,
        'valign': 'vcenter',
        'align': 'center',
        'text_wrap': True,
        'border': 1
    })
    
    belum_dikonfirmasi_date_format = workbook.add_format({
        'font_size': 22,
        'font_color': '#000000',  # Black text (changed from white)
        'bg_color': '#FF0000',  # Red background
        'bold': True,
        'valign': 'vcenter',
        'align': 'center',
        'text_wrap': True,
        'border': 1,
        'num_format': 'dd/mm/yyyy'
    })
    
    # Set page layout
    worksheet.set_margins(left=0.28, right=0.28, top=0.75, bottom=0.75)
    worksheet.set_landscape()
    worksheet.hide_gridlines(2)
    worksheet.fit_to_pages(1, 0)  # Fit to 1 page wide, unlimited pages tall
    worksheet.set_zoom(70)  # Zoom ke 70% agar lebih rapi dilihat
    
    # Set column widths dengan text wrap
    col_widths = {
        'A': 8,      # Nomor
        'B': 16,     # Cabang
        'C': 16,     # Tanggal Kirim
        'D': 16,     # Tanggal Potong
        'E': 20,     # No. Invoice
        'F': 22,     # Status Perkembangan
        'G': 18,     # Pemotongan Real
        'H': 20,     # Nama Anak
        'I': 18,     # Jenis Kelamin Anak
        'J': 20,     # Nama Bapak
        'K': 18,     # Telpon 1
        'L': 18,     # Telpon 2
        'M': 20,     # Paket
        'N': 14,     # Jumlah
        'O': 24,     # Menu
        'P': 24,     # Pemotongan Disaksikan
        'Q': 35,     # Catatan Khusus
        'R': 12      # CS
    }
    
    for col, width in col_widths.items():
        worksheet.set_column(f'{col}:{col}', width)
    
    # PERTAMA: Write header SEBELUM data dengan font size 22
    for col_num, value in enumerate(df.columns.values):
        worksheet.write(0, col_num, value, header_format)
    # Set auto height untuk header dengan minimum 35
    worksheet.set_row(0, None)  # Auto height
    
    # KEDUA: FORMAT data rows yang sudah ditulis to_excel dengan format yang sesuai
    # Data sudah ada di row 1+ dari to_excel, kita tinggal apply format saja
    for row_num in range(1, len(df) + 1):
        worksheet.set_row(row_num, None)  # Auto height untuk data rows
        
        # Cek status perkembangan untuk baris ini
        row_data = df.iloc[row_num - 1]
        status_perkembangan = str(row_data['Status Perkembangan']).strip() if 'Status Perkembangan' in df.columns else ''
        is_belum_dikonfirmasi = 'belum dikonfirmasi' in status_perkembangan.lower()
        
        # Format setiap cell di row ini
        for col_num, cell_value in enumerate(df.iloc[row_num - 1].values):
            # Default format
            fmt = cell_format
            col_name = df.columns[col_num]
            
            # Handle NaT values and convert to proper types for xlsxwriter
            display_value = cell_value
            
            # Convert pandas/numpy types to Python native types
            if pd.isna(cell_value):
                display_value = ''
            elif isinstance(cell_value, (pd.Timestamp, datetime.datetime)):
                display_value = cell_value.to_pydatetime() if hasattr(cell_value, 'to_pydatetime') else cell_value
            elif isinstance(cell_value, (np.integer, np.floating)):
                display_value = cell_value.item()
            elif isinstance(cell_value, np.bool_):
                display_value = bool(cell_value)
            elif isinstance(cell_value, (np.ndarray, pd.Series)):
                display_value = str(cell_value)
            elif hasattr(cell_value, 'item'):  # other numpy types
                try:
                    display_value = cell_value.item()
                except:
                    display_value = str(cell_value) if cell_value is not None else ''
            
            # Safe string conversion for cell_str
            cell_str = str(cell_value).strip() if cell_value is not None and not pd.isna(cell_value) else ''
            
            # LOGIKA BARU: Jika "Belum Dikonfirmasi", set default ke format merah dengan font hitam
            # Tapi jika ada formatting khusus di kolom tertentu, formatting khusus menimpa background merah
            if is_belum_dikonfirmasi:
                # Default untuk baris "Belum Dikonfirmasi" adalah background merah dengan font hitam
                if col_name in ['Tanggal Kirim', 'Tanggal Potong']:
                    fmt = belum_dikonfirmasi_date_format
                elif col_name in ['Nomor', 'Jumlah', 'Jenis Kelamin Anak']:
                    fmt = belum_dikonfirmasi_center_format
                else:
                    fmt = belum_dikonfirmasi_format
                
                # Sekarang cek apakah ada formatting khusus yang harus menimpa background merah
                # Formatting khusus akan diterapkan di bawah, jadi kita perlu lanjut ke pengecekan
            
            # === SEKARANG CEK FORMATTING KHUSUS - ini akan menimpa format merah jika ada ===
            
            # === KOLOM C & D: Cek jika Tanggal Kirim = Tanggal Potong ===
            if col_name == 'Tanggal Kirim':
                # Bandingkan dengan Tanggal Potong (kolom D)
                tanggal_potong_val = row_data[3]
                tanggal_kirim_val = cell_value
                
                # Handle datetime objects - extract date part only
                try:
                    if pd.notna(tanggal_kirim_val) and pd.notna(tanggal_potong_val):
                        tanggal_kirim_str = pd.to_datetime(tanggal_kirim_val).strftime('%Y-%m-%d')
                        tanggal_potong_str = pd.to_datetime(tanggal_potong_val).strftime('%Y-%m-%d')
                        if tanggal_kirim_str == tanggal_potong_str:
                            # Formatting khusus: tanggal sama menimpa format merah
                            fmt = date_same_format
                        elif not is_belum_dikonfirmasi:
                            # Hanya set date_format jika bukan "Belum Dikonfirmasi"
                            fmt = date_format
                    elif not is_belum_dikonfirmasi:
                        fmt = cell_center_format
                except:
                    if not is_belum_dikonfirmasi:
                        fmt = cell_center_format
            
            elif col_name == 'Tanggal Potong':
                # Bandingkan dengan Tanggal Kirim (kolom C)
                tanggal_kirim_val = row_data[2]
                tanggal_potong_val = cell_value
                
                # Handle datetime objects - extract date part only
                try:
                    if pd.notna(tanggal_kirim_val) and pd.notna(tanggal_potong_val):
                        tanggal_kirim_str = pd.to_datetime(tanggal_kirim_val).strftime('%Y-%m-%d')
                        tanggal_potong_str = pd.to_datetime(tanggal_potong_val).strftime('%Y-%m-%d')
                        if tanggal_kirim_str == tanggal_potong_str:
                            # Formatting khusus: tanggal sama menimpa format merah
                            fmt = date_same_format
                        elif not is_belum_dikonfirmasi:
                            # Hanya set date_format jika bukan "Belum Dikonfirmasi"
                            fmt = date_format
                    elif not is_belum_dikonfirmasi:
                        fmt = cell_center_format
                except:
                    if not is_belum_dikonfirmasi:
                        fmt = cell_center_format
            
            # === KOLOM M: Paket - Cek Jantan atau Kebuli ===
            elif col_name == 'Paket':
                if 'jantan' in cell_str.lower():
                    # Formatting khusus: paket jantan menimpa format merah
                    fmt = paket_jantan_format
                elif 'kebuli' in cell_str.lower():
                    # Formatting khusus: paket kebuli menimpa format merah
                    fmt = paket_kebuli_format
                elif not is_belum_dikonfirmasi:
                    fmt = cell_format
            
            # === KOLOM P: Pemotongan Disaksikan - Cek Live Video Call atau Disaksikan ===
            elif col_name == 'Pemotongan Disaksikan':
                if 'live video call' in cell_str.lower():
                    # Formatting khusus: live video call menimpa format merah
                    fmt = pemotongan_live_format
                elif 'disaksikan' in cell_str.lower():
                    # Formatting khusus: disaksikan menimpa format merah
                    fmt = pemotongan_disaksikan_format
                elif not is_belum_dikonfirmasi:
                    fmt = cell_format
            
            # === KOLOM Q: Catatan Khusus - Cek Domba, Kambing, Upgrade Bobot, atau Bukan Aqiqah ===
            elif col_name == 'Catatan Khusus':
                keywords = ['domba', 'kambing', 'upgrade bobot', 'bukan aqiqah']
                if any(keyword in cell_str.lower() for keyword in keywords):
                    # Formatting khusus: catatan khusus menimpa format merah
                    fmt = catatan_yellow_format
                elif not is_belum_dikonfirmasi:
                    fmt = cell_format
            
            # === KOLOM Nomor, Jumlah, Jenis Kelamin Anak - Center Aligned ===
            elif col_name in ['Nomor', 'Jumlah', 'Jenis Kelamin Anak']:
                if not is_belum_dikonfirmasi:
                    fmt = cell_center_format
            
            # === Apply date num_format to date columns ===
            if col_name in ['Tanggal Kirim', 'Tanggal Potong'] and fmt == date_format:
                # Already has num_format applied in date_format definition
                pass
            
            worksheet.write(row_num, col_num, display_value, fmt)

# ============================================================
# HELPER FUNCTIONS: KATEGORI MANAGEMENT
# ============================================================
KATEGORI_FILE = "kategori.csv"

def load_kategori():
    """Load kategori from CSV file, return empty df if not exists"""
    if os.path.exists(KATEGORI_FILE):
        return pd.read_csv(KATEGORI_FILE)
    return pd.DataFrame(columns=['Nama Barang', 'Kategori Final'])

def save_kategori(df):
    """Save kategori to CSV file"""
    df.to_csv(KATEGORI_FILE, index=False)
    st.success("✅ Kategori berhasil disimpan!")

def add_kategori(nama_barang, kategori_final):
    """Add new kategori entry"""
    df_kat = load_kategori()
    new_row = pd.DataFrame({'Nama Barang': [nama_barang], 'Kategori Final': [kategori_final]})
    df_kat = pd.concat([df_kat, new_row], ignore_index=True)
    save_kategori(df_kat)
    return df_kat

def delete_kategori(idx):
    """Delete kategori by index"""
    df_kat = load_kategori()
    df_kat = df_kat.drop(idx).reset_index(drop=True)
    save_kategori(df_kat)
    return df_kat

def update_kategori(idx, nama_barang, kategori_final):
    """Update kategori entry"""
    df_kat = load_kategori()
    df_kat.at[idx, 'Nama Barang'] = nama_barang
    df_kat.at[idx, 'Kategori Final'] = kategori_final
    save_kategori(df_kat)
    return df_kat

# =============================================================================
# FUNGSI 1: TRANSFORMASI REKAP PEMOTONGAN
# =============================================================================
def transform_rekap_pemotongan(uploaded_file):
    try:
        if uploaded_file.name.endswith('.xlsx'):
            df = read_excel_robust(uploaded_file, header=1)
        elif uploaded_file.name.endswith('.csv'):
            df = pd.read_csv(uploaded_file, header=1)
        else:
            st.error("Format file tidak didukung. Harap unggah file .xlsx atau .csv")
            return None
    except Exception as e:
        st.error(f"Gagal membaca file: {e}")
        st.warning("Tips: Jika file .xlsx Anda error, coba simpan sheet pertama sebagai file CSV dan unggah kembali file CSV tersebut.")
        return None

    df.dropna(subset=['Cabang'], inplace=True)
    df.dropna(subset=['Tanggal Kirim'], inplace=True)
    df = df[df['Cabang'] != 'Cabang'].copy()
    
    if "Jenis Kelamin AnakNama" in df.columns:
        df.rename(columns={"Jenis Kelamin AnakNama": "Jenis Kelamin Anak"}, inplace=True)
    if "Pemotongan DisaksikanNama" in df.columns:
        df.rename(columns={"Pemotongan DisaksikanNama": "Pemotongan Disaksikan"}, inplace=True)
    
    # Cek berbagai variasi nama kolom untuk Pemotongan Disaksikan
    possible_names = [col for col in df.columns if 'pemotongan' in col.lower() and 'disaksikan' in col.lower()]
    if possible_names and 'Pemotongan Disaksikan' not in df.columns:
        df.rename(columns={possible_names[0]: 'Pemotongan Disaksikan'}, inplace=True)
    
    # Parse tanggal - coba format DD/MM/YYYY dulu, jika gagal coba format lain
    def parse_tanggal(val):
        if pd.isna(val):
            return None
        try:
            # Coba format DD/MM/YYYY (asumsi input user)
            return pd.to_datetime(str(val).strip(), format='%d/%m/%Y')
        except:
            try:
                # Coba format lain (auto-detect)
                result = pd.to_datetime(str(val).strip(), errors='coerce')
                return result
            except:
                return None
    
    df['Tanggal Kirim'] = df['Tanggal Kirim'].apply(parse_tanggal)
    df['Tanggal Potong'] = df['Tanggal Potong'].apply(parse_tanggal)
    
    for col in ['Telpon 1', 'Telpon 2']:
        if col in df.columns:
            df[col] = df[col].astype(str).str.replace(r'\.0$', '', regex=True).str.replace(r'\D', '', regex=True)

    df_menu_prep = df[['No. Invoice', 'Paket & Menu', 'Jumlah']].copy()
    df_menu_prep = df_menu_prep[~df_menu_prep['Paket & Menu'].astype(str).str.contains("Paket", na=False)]
    df_menu_prep['Menu_Item'] = df_menu_prep['Paket & Menu'].astype(str) + " " + df_menu_prep['Jumlah'].astype(str)
    df_menu_final = df_menu_prep.groupby('No. Invoice')['Menu_Item'].apply(lambda x: ', '.join(x) + ' PORSI').reset_index()
    df_menu_final.rename(columns={'Menu_Item': 'Menu'}, inplace=True)

    # Agregasi Jumlah berdasarkan No. Invoice + Paket & Menu
    df_paket_prep = df[['No. Invoice', 'Paket & Menu', 'Jumlah']].copy()
    df_paket_prep = df_paket_prep[df_paket_prep['Paket & Menu'].astype(str).str.contains("Paket", na=False)].copy()
    
    # Convert Jumlah to numeric untuk penjumlahan
    df_paket_prep['Jumlah'] = pd.to_numeric(df_paket_prep['Jumlah'], errors='coerce').fillna(0).astype(int)
    
    # Aggregate berdasarkan KEDUANYA (No. Invoice + Paket & Menu)
    # Jika No. Invoice sama & Paket sama → jumlahkan Jumlah menjadi 1 row
    # Jika No. Invoice sama tapi Paket berbeda → tetap terpisah
    df_paket_final = df_paket_prep.groupby(['No. Invoice', 'Paket & Menu']).agg({
        'Jumlah': 'sum'
    }).reset_index()
    
    # Jika ada multiple Paket per Invoice, ambil yang pertama (biasanya hanya 1 Paket per Invoice)
    # Tapi jika ada 2 Paket berbeda, tetap terpisah
    df_paket_final = df_paket_final.sort_values('No. Invoice').reset_index(drop=True)
    df_paket_final.rename(columns={'Paket & Menu': 'Paket'}, inplace=True)
    
    cols_to_drop = ['Paket & Menu', 'No. Urut', 'No. Domba', 'Satuan',
                    'Tanggal Domba Dipotong', 'Jam Tiba (hh:mm)', 'Jam Kirim (hh:mm)', 'Kode Menu']
    df_base = df.drop(columns=cols_to_drop + ['Jumlah'], errors='ignore')
    
    # Custom aggregation untuk Status Perkembangan
    # Prioritas: ambil "Belum Dikonfirmasi" jika ada, jika tidak ambil yang pertama
    def aggregate_status(statuses):
        if any('belum dikonfirmasi' in str(s).lower() for s in statuses):
            return [s for s in statuses if 'belum dikonfirmasi' in str(s).lower()][0]
        return statuses.iloc[0]
    
    # Untuk agregasi yang benar, kita perlu akses ke kolom 'Paket & Menu' (kolom Q)
    # Jadi kita tidak drop 'Paket & Menu' dari df_base, kita simpan untuk referensi
    # Buat df yang include 'Paket & Menu' untuk agregasi custom
    df_with_paket = df.copy()
    
    # Custom aggregation untuk kolom yang perlu mengikuti baris dengan "Paket" di kolom Q
    def aggregate_by_paket(group_df, col_name):
        # Cari baris yang di kolom 'Paket & Menu' mengandung kata "Paket"
        paket_rows = group_df[group_df['Paket & Menu'].astype(str).str.contains("Paket", case=False, na=False)]
        
        if len(paket_rows) > 0:
            # Ambil nilai dari baris yang ada "Paket" (ambil yang pertama jika ada multiple)
            return paket_rows.iloc[0][col_name]
        else:
            # Jika tidak ada baris dengan "Paket", ambil nilai pertama
            return group_df.iloc[0][col_name]
    
    # Aggregate data berdasarkan No. Invoice dengan logika khusus
    grouped = df_with_paket.groupby('No. Invoice')
    
    result_rows = []
    for invoice, group in grouped:
        result_row = {'No. Invoice': invoice}
        
        for col in df_base.columns:
            if col == 'No. Invoice':
                continue
            elif col == 'Status Perkembangan':
                result_row[col] = aggregate_status(group[col])
            elif col == 'Pemotongan Disaksikan':
                # Gunakan agregasi berdasarkan baris yang punya "Paket"
                result_row[col] = aggregate_by_paket(group, col)
            else:
                # Untuk kolom lainnya, juga ikuti baris dengan "Paket" jika ada
                result_row[col] = aggregate_by_paket(group, col)
        
        result_rows.append(result_row)
    
    df_base_aggregated = pd.DataFrame(result_rows)

    # Merge dengan paket (bisa multiple rows jika ada berbeda paket)
    df_merged = pd.merge(df_base_aggregated, df_paket_final, on='No. Invoice', how='left')
    df_merged = pd.merge(df_merged, df_menu_final[['No. Invoice', 'Menu']], on='No. Invoice', how='left')

    df_final = df_merged.copy()
    df_final['Pemotongan Real'] = ''

    final_columns_order = [
        "Cabang", "Tanggal Kirim", "Tanggal Potong", "No. Invoice", "Status Perkembangan", 
        "Pemotongan Real", "Nama Anak", "Jenis Kelamin Anak", "Nama Bapak", 
        "Telpon 1", "Telpon 2", "Paket", "Jumlah", "Menu", "Pemotongan Disaksikan", 
        "Catatan Khusus", "CS"
    ]
    
    existing_columns = [col for col in final_columns_order if col in df_final.columns]
    df_final = df_final[existing_columns]
    df_final.insert(0, 'Nomor', range(1, len(df_final) + 1))
    
    return df_final

# =============================================================================
# FUNGSI 2: TRANSFORMASI REKAP KEBUTUHAN MINGGUAN
# =============================================================================
def transform_rekap_kebutuhan(file_sales):
    try:
        df_sales = read_excel_robust(file_sales, header=1)
    except Exception as e:
        st.error(f"Gagal membaca file: {e}")
        st.warning("Pastikan file Excel Anda valid dan berisi data yang sesuai.")
        return None

    # Load kategori from external file
    df_kategori = load_kategori()
    
    if len(df_kategori) == 0:
        st.error("❌ Kategori kosong! Silakan tambahkan kategori terlebih dahulu.")
        return None
    
    # Clean up data
    df_sales.dropna(subset=['Tanggal Potong', 'No. Invoice'], inplace=True)
    df_sales = df_sales[df_sales['Cabang'] != 'Cabang'].copy()
    
    # Parse dates dengan format DD/MM/YYYY
    def parse_tanggal_kebutuhan(val):
        if pd.isna(val):
            return None
        try:
            # Coba format DD/MM/YYYY (asumsi input user)
            return pd.to_datetime(str(val).strip(), format='%d/%m/%Y')
        except:
            try:
                # Coba format lain (auto-detect)
                result = pd.to_datetime(str(val).strip(), errors='coerce')
                return result
            except:
                return None
    
    df_sales['Tanggal Potong'] = df_sales['Tanggal Potong'].apply(parse_tanggal_kebutuhan)
    
    # Get Paket & Menu from column Q (index 16) and Jumlah from column R (index 17)
    try:
        # Get the column names (Excel might have them or might not)
        if 'Paket & Menu' not in df_sales.columns:
            # Try by column index
            paket_col = df_sales.iloc[:, 16]
            jumlah_col = df_sales.iloc[:, 17]
        else:
            paket_col = df_sales['Paket & Menu']
            jumlah_col = df_sales['Jumlah']
        
        df_sales['Paket & Menu'] = paket_col
        df_sales['Jumlah'] = pd.to_numeric(jumlah_col, errors='coerce').fillna(0).astype(int)
    except Exception as e:
        st.error(f"Tidak dapat menemukan data paket & menu atau jumlah: {e}")
        return None
    
    # Groupby Tanggal Potong and Paket & Menu, aggregating Jumlah
    df_grouped = df_sales.groupby(["Tanggal Potong", "Paket & Menu"])['Jumlah'].sum().reset_index()
    df_grouped['Jumlah'] = df_grouped['Jumlah'].astype(int)
    
    # Format dates to string for pivot table
    df_grouped['Tanggal Potong_str'] = df_grouped['Tanggal Potong'].apply(
        lambda x: pd.to_datetime(x).strftime('%d-%m-%Y')
    )
    
    # Create pivot table with ALL data first
    df_pivot = df_grouped.pivot_table(
        index='Paket & Menu',
        columns='Tanggal Potong_str',
        values='Jumlah',
        aggfunc='sum'
    ).fillna(0).astype(int)
    
    df_pivot.reset_index(inplace=True)
    
    # Separate items that are in kategori vs not in kategori
    kategori_items = df_kategori['Nama Barang'].values
    df_in_kategori = df_pivot[df_pivot['Paket & Menu'].isin(kategori_items)].copy()
    df_not_in_kategori = df_pivot[~df_pivot['Paket & Menu'].isin(kategori_items)].copy()
    
    # Store missing items for display
    st.session_state.missing_items = df_not_in_kategori if len(df_not_in_kategori) > 0 else None
    
    # Use only items in kategori for main output
    df_pivot = df_in_kategori.copy()
    
    # Sort date columns
    date_cols = [c for c in df_pivot.columns if c != 'Paket & Menu']
    date_cols_sorted = sorted(date_cols, key=lambda x: pd.to_datetime(x, format='%d-%m-%Y'))
    
    # Filter hanya 7 hari: hari ini + 6 hari berikutnya
    today = pd.Timestamp.now().normalize()
    target_dates = [today + pd.Timedelta(days=i) for i in range(7)]
    target_dates_str = [d.strftime('%d-%m-%Y') for d in target_dates]
    
    # Ambil hanya kolom tanggal yang ada dalam 7 hari target
    date_cols_filtered = [d for d in date_cols_sorted if d in target_dates_str]
    
    # Add weekday names to date columns
    hari_id = ['Senin', 'Selasa', 'Rabu', 'Kamis', 'Jumat', 'Sabtu', 'Minggu']
    date_cols_renamed = {}
    for date_str in date_cols_filtered:
        date_obj = pd.to_datetime(date_str, format='%d-%m-%Y')
        weekday_name = hari_id[date_obj.weekday()]
        new_col_name = f"{weekday_name}\n{date_str}"
        date_cols_renamed[date_str] = new_col_name
    
    df_pivot.rename(columns=date_cols_renamed, inplace=True)
    date_cols_sorted_renamed = [date_cols_renamed[c] for c in date_cols_filtered]
    
    # Reorder columns - hanya ambil kolom yang ada dalam 7 hari
    final_col_order = ['Paket & Menu'] + date_cols_sorted_renamed
    df_final = df_pivot[final_col_order].copy()
    
    # Add kolom Total untuk setiap baris (jumlah dari semua tanggal)
    if len(date_cols_sorted_renamed) > 0:
        df_final['Total'] = df_final[date_cols_sorted_renamed].sum(axis=1).astype(int)
    
    # Filter: hapus baris yang Total = 0
    df_final = df_final[df_final['Total'] > 0].copy()
    
    # Add TOTAL row at the bottom
    if len(date_cols_sorted_renamed) > 0:
        totals_row = {'Paket & Menu': 'TOTAL'}
        for col in date_cols_sorted_renamed:
            totals_row[col] = int(df_final[col].sum())
        # Total dari kolom Total
        totals_row['Total'] = int(df_final['Total'].sum())
        df_final = pd.concat([df_final, pd.DataFrame([totals_row])], ignore_index=True)
    
    return df_final

# =============================================================================
# FUNGSI 3: TRANSFORMASI LABEL MASAK
# =============================================================================
def transform_and_create_word_label(file_input):
    """
    Membaca file Excel, mentransformasikannya, dan menghasilkan file Word.
    """
    try:
        # Baca sheet pertama
        df = read_excel_robust(file_input, header=1)
        
        df.dropna(subset=['No. Invoice'], inplace=True)
        df = df[df['Cabang'] != 'Cabang'].copy()

        # Parse tanggal dengan format DD/MM/YYYY
        def parse_tanggal_label(val):
            if pd.isna(val):
                return None
            try:
                # Coba format DD/MM/YYYY (asumsi input user)
                return pd.to_datetime(str(val).strip(), format='%d/%m/%Y')
            except:
                try:
                    # Coba format lain (auto-detect)
                    result = pd.to_datetime(str(val).strip(), errors='coerce')
                    return result
                except:
                    return None
        
        for col in ['Tanggal Kirim', 'Tanggal Potong']:
            df[col] = df[col].apply(parse_tanggal_label)
            df[col] = df[col].dt.strftime('%d/%m/%Y')
        
        def format_time(time_val):
            if isinstance(time_val, datetime.time):
                return time_val.strftime('%H:%M')
            elif isinstance(time_val, str):
                return time_val
            return str(time_val)

        df['Jam Tiba (hh:mm)'] = df['Jam Tiba (hh:mm)'].apply(format_time)
        df['Jam Kirim (hh:mm)'] = df['Jam Kirim (hh:mm)'].apply(format_time)

        df.fillna('', inplace=True)
        
        df['Detail Customer'] = (
            "Nama Aqiqah:\n" + df['Nama Anak'].astype(str).str.strip() + "\n" +
            "No. Invoice: " + df['No. Invoice'].astype(str).str.strip() + "\n" +
            "Jenis Kelamin: " + df['Jenis Kelamin Anak'].astype(str).str.strip() + "\n" +
            "Cabang: " + df['Cabang'].astype(str).str.strip()
        )
        df['Detail Waktu'] = (
            "Tgl Kirim: " + df['Tanggal Kirim'].astype(str).str.strip() + "\n" +
            "Tgl Potong: " + df['Tanggal Potong'].astype(str).str.strip() + "\n" +
            "Jam Tiba: " + df['Jam Tiba (hh:mm)'].astype(str).str.strip() + "\n" +
            "Jam Kirim: " + df['Jam Kirim (hh:mm)'].astype(str).str.strip()
        )
        # Format Menu: remove decimal points, keep only integers
        df['Menu'] = df['Paket & Menu'].astype(str) + " " + df['Jumlah'].astype(str).str.replace(r'\.0$', '', regex=True) + " " + df['Satuan'].astype(str)
        df['Berat'] = "Berat |\n....... KG"

        df_final = df[['Detail Customer', 'Detail Waktu', 'Menu', 'Berat', 'Cabang']].copy()
        
        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(0.88)
            section.bottom_margin = Cm(1.75)
            section.left_margin = Cm(2.12)
            section.right_margin = Cm(1.42)

        for i in range(0, len(df_final), 5):
            chunk = df_final.iloc[i:i+5]
            
            table = doc.add_table(rows=0, cols=4)
            table.style = 'Table Grid'

            # Tambahkan baris data (tanpa header)
            for df_index, record in chunk.iterrows():
                row_cells = table.add_row().cells
                
                # Kolom 1: Detail Customer dengan formatting khusus
                cell_col1 = row_cells[0]
                cell_col1.text = ''
                detail_customer = str(record['Detail Customer'])
                cabang = str(record['Cabang']).strip()
                
                for line in detail_customer.split('\n'):
                    if cell_col1.text == '':
                        p = cell_col1.paragraphs[0]
                    else:
                        p = cell_col1.add_paragraph()
                    
                    if line.startswith('Nama Aqiqah:'):
                        # Bold untuk "Nama Aqiqah:" saja (nama ada di baris berikutnya)
                        run_label = p.add_run(line)
                        run_label.font.bold = True
                        run_label.font.name = 'Arial'
                        run_label.font.size = Pt(10)
                    elif line and not line.startswith('No. Invoice:') and not line.startswith('Jenis Kelamin:') and not line.startswith('Cabang:'):
                        # Ini adalah nama anak (baris setelah "Nama Aqiqah:")
                        run_value = p.add_run(line)
                        run_value.font.bold = True
                        run_value.font.name = 'Arial'
                        run_value.font.size = Pt(10)
                    elif line.startswith('Cabang:'):
                        # Cabang dengan warna berdasarkan value
                        run_label = p.add_run('Cabang: ')
                        run_label.font.name = 'Arial'
                        run_label.font.size = Pt(10)
                        
                        cabang_value = line.replace('Cabang: ', '')
                        run_value = p.add_run(cabang_value)
                        run_value.font.name = 'Arial'
                        run_value.font.size = Pt(10)
                        
                        if 'Cibubur' in cabang_value:
                            run_value.font.bold = True
                            run_value.font.color.rgb = RGBColor(0, 0, 255)
                        elif 'JaDeTa' in cabang_value:
                            run_value.font.bold = True
                            run_value.font.color.rgb = RGBColor(255, 0, 0)
                    else:
                        # Normal text untuk line lainnya
                        run = p.add_run(line)
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                
                # Kolom 2: Detail Waktu dengan Tgl Kirim di bold
                cell_col2 = row_cells[1]
                cell_col2.text = ''
                detail_waktu = str(record['Detail Waktu'])
                
                for line in detail_waktu.split('\n'):
                    if cell_col2.text == '':
                        p = cell_col2.paragraphs[0]
                    else:
                        p = cell_col2.add_paragraph()
                    
                    if line.startswith('Tgl Kirim:'):
                        # Bold untuk "Tgl Kirim: [value]"
                        run_label = p.add_run('Tgl Kirim: ')
                        run_label.font.bold = True
                        run_label.font.name = 'Arial'
                        run_label.font.size = Pt(10)
                        
                        run_value = p.add_run(line.replace('Tgl Kirim: ', ''))
                        run_value.font.bold = True
                        run_value.font.name = 'Arial'
                        run_value.font.size = Pt(10)
                    else:
                        # Normal text untuk line lainnya
                        run = p.add_run(line)
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                
                # Kolom 3: Menu di bold semua
                cell_col3 = row_cells[2]
                cell_col3.text = ''
                menu_text = str(record['Menu'])
                
                p = cell_col3.paragraphs[0]
                run = p.add_run(menu_text)
                run.font.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                
                # Kolom 4: Berat normal
                cell_col4 = row_cells[3]
                cell_col4.text = str(record['Berat'])

            # Atur format untuk seluruh tabel
            for row in table.rows:
                row.height = Cm(4.5)

            # Atur lebar kolom
            table.columns[0].width = Cm(5.64)
            table.columns[1].width = Cm(2.75)
            table.columns[2].width = Cm(5.29)
            table.columns[3].width = Cm(1.76)

            if i + 5 < len(df_final):
                doc.add_page_break()

        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io

    except Exception as e:
        st.error(f"Terjadi kesalahan: {e}")
        st.warning("Pastikan file yang diunggah adalah template yang benar dan memiliki struktur data yang sesuai.")
        return None

# =============================================================================
# BAGIAN TAMPILAN APLIKASI WEB (USER INTERFACE)
# =============================================================================

# Set page config untuk layout yang lebih baik
st.set_page_config(
    page_title="Riyadh Aqiqah - Management System",
    page_icon="🐑",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS untuk membuat tampilan lebih modern dan clean
st.markdown("""
<style>
    /* Root background - area di luar konten utama */
    .stApp {
        background: linear-gradient(135deg, #065f46 0%, #047857 50%, #059669 100%);
    }
    
    /* Main background */
    .main {
        background: linear-gradient(135deg, #f0fdf4 0%, #dcfce7 100%);
    }
    
    /* Sidebar styling */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #059669 0%, #047857 100%);
        padding-top: 2rem;
    }
    
    [data-testid="stSidebar"] .css-1d391kg, [data-testid="stSidebar"] .st-emotion-cache-1d391kg {
        color: white;
    }
    
    /* Sidebar title */
    [data-testid="stSidebar"] h1 {
        color: white !important;
        font-weight: 700;
        padding: 1rem 0;
        border-bottom: 2px solid rgba(255, 255, 255, 0.2);
        margin-bottom: 1.5rem;
    }
    
    /* Radio buttons in sidebar */
    [data-testid="stSidebar"] label {
        color: white !important;
        font-size: 1rem !important;
        font-weight: 500 !important;
    }
    
    [data-testid="stSidebar"] [role="radiogroup"] label {
        padding: 0.75rem 1rem;
        border-radius: 8px;
        transition: all 0.3s ease;
        margin: 0.25rem 0;
    }
    
    [data-testid="stSidebar"] [role="radiogroup"] label:hover {
        background: rgba(255, 255, 255, 0.1);
    }
    
    /* Card-like containers */
    .stAlert, .element-container {
        border-radius: 12px;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.1);
    }
    
    /* Main title styling */
    h1 {
        color: #065f46;
        font-weight: 700;
        padding-bottom: 1rem;
        border-bottom: 3px solid #10b981;
        margin-bottom: 2rem;
    }
    
    /* Buttons styling */
    .stButton > button {
        background: linear-gradient(135deg, #10b981 0%, #059669 100%);
        color: white;
        font-weight: 600;
        border: none;
        border-radius: 8px;
        padding: 0.5rem 2rem;
        transition: all 0.3s ease;
        box-shadow: 0 4px 12px rgba(16, 185, 129, 0.3);
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 16px rgba(16, 185, 129, 0.4);
    }
    
    /* Download button special styling */
    .stDownloadButton > button {
        background: linear-gradient(135deg, #11998e 0%, #38ef7d 100%);
        color: white;
        font-weight: 600;
        border: none;
        border-radius: 8px;
        padding: 0.5rem 2rem;
        box-shadow: 0 4px 12px rgba(17, 153, 142, 0.3);
    }
    
    .stDownloadButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 16px rgba(17, 153, 142, 0.4);
    }
    
    /* File uploader styling */
    [data-testid="stFileUploader"] {
        background: rgba(255, 255, 255, 0.7);
        padding: 2rem;
        border-radius: 12px;
        border: 2px dashed #6ee7b7;
        box-shadow: 0 2px 8px rgba(16, 185, 129, 0.1);
        transition: all 0.3s ease;
    }
    
    [data-testid="stFileUploader"]:hover {
        background: rgba(255, 255, 255, 0.9);
        border-color: #10b981;
        box-shadow: 0 4px 12px rgba(16, 185, 129, 0.25);
    }
    
    /* File uploader label text */
    [data-testid="stFileUploader"] label {
        color: #1f2937 !important;
        font-weight: 600 !important;
    }
    
    [data-testid="stFileUploader"] small {
        color: #374151 !important;
    }
    
    /* Success/Info/Warning messages */
    .stSuccess {
        background: linear-gradient(135deg, #d1fae5 0%, #a7f3d0 100%);
        color: #065f46;
        border-radius: 12px;
        padding: 1rem;
        border-left: 4px solid #10b981;
    }
    
    .stInfo {
        background: linear-gradient(135deg, #e0f2fe 0%, #bae6fd 100%);
        color: #075985;
        border-radius: 12px;
        padding: 1rem;
        border-left: 4px solid #0284c7;
    }
    
    .stWarning {
        background: linear-gradient(135deg, #fef3c7 0%, #fde68a 100%);
        color: #92400e;
        border-radius: 12px;
        padding: 1rem;
        border-left: 4px solid #f59e0b;
    }
    
    /* Dataframe styling */
    [data-testid="stDataFrame"] {
        border-radius: 12px;
        overflow: hidden;
        box-shadow: 0 4px 16px rgba(0, 0, 0, 0.1);
    }
    
    /* Expander styling */
    .streamlit-expanderHeader {
        background: white;
        border-radius: 8px;
        font-weight: 600;
        color: #065f46;
    }
    
    /* Tabs styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    
    .stTabs [data-baseweb="tab"] {
        background: white;
        border-radius: 8px 8px 0 0;
        padding: 0.75rem 1.5rem;
        font-weight: 600;
        border: 2px solid transparent;
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #10b981 0%, #059669 100%);
        color: white !important;
    }
    
    /* Input fields */
    .stTextInput > div > div > input {
        border-radius: 8px;
        border: 2px solid #e2e8f0;
        padding: 0.75rem;
        transition: all 0.3s ease;
    }
    
    .stTextInput > div > div > input:focus {
        border-color: #10b981;
        box-shadow: 0 0 0 3px rgba(16, 185, 129, 0.1);
    }
    
    /* Selectbox */
    .stSelectbox > div > div {
        border-radius: 8px;
        border: 2px solid #e2e8f0;
    }
    
    /* Divider */
    hr {
        margin: 2rem 0;
        border: none;
        height: 2px;
        background: linear-gradient(90deg, transparent, #cbd5e0, transparent);
    }
    
    /* Animation for page load */
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(20px); }
        to { opacity: 1; transform: translateY(0); }
    }
    
    .element-container {
        animation: fadeIn 0.5s ease-out;
    }
</style>
""", unsafe_allow_html=True)

# Header dengan logo dan deskripsi
st.markdown("""
<div style='text-align: center; padding: 2rem 0; background: white; border-radius: 16px; 
            box-shadow: 0 4px 16px rgba(0,0,0,0.1); margin-bottom: 2rem;'>
    <img src='https://down-id.img.susercontent.com/file/id-11134004-7r98y-lmmve5folrlzb8' 
         style='width: 120px; height: 120px; margin-bottom: 1rem; border-radius: 12px;
                box-shadow: 0 4px 12px rgba(16, 185, 129, 0.2);' />
    <h1 style='color: #065f46; font-size: 2.5rem; margin: 0; border: none;'>
        Riyadh Aqiqah Management System
    </h1>
    <p style='color: #64748b; font-size: 1.1rem; margin-top: 0.5rem;'>
        Sistem manajemen data pemotongan, kebutuhan mingguan, dan label masak
    </p>
</div>
""", unsafe_allow_html=True)

# Sidebar with better styling
st.sidebar.markdown("""
<div style='text-align: center; padding: 1rem 0; margin-bottom: 1rem;'>
    <img src='https://down-id.img.susercontent.com/file/id-11134004-7r98y-lmmve5folrlzb8' 
         style='width: 80px; height: 80px; border-radius: 12px; margin-bottom: 0.5rem;
                box-shadow: 0 4px 12px rgba(0, 0, 0, 0.2);' />
    <p style='color: white; margin: 0.5rem 0 0 0; font-size: 1rem; font-weight: 600;'>
        Riyadh Aqiqah
    </p>
</div>
""", unsafe_allow_html=True)

st.sidebar.title("📋 Navigasi Menu")
st.sidebar.markdown("<br>", unsafe_allow_html=True)

menu_pilihan = st.sidebar.radio(
    "Pilih menu yang ingin Anda gunakan:",
    ("Rekap Pemotongan", "Rekap Kebutuhan Mingguan", "Label Masak"),
    label_visibility="collapsed"
)

# Info footer di sidebar
st.sidebar.markdown("<br><br>", unsafe_allow_html=True)
st.sidebar.markdown("""
<div style='padding: 1rem; margin-top: 2rem;
            background: rgba(255,255,255,0.1); border-radius: 8px;'>
    <p style='color: rgba(255,255,255,0.9); font-size: 0.85rem; margin: 0; line-height: 1.5;'>
        💡 <b>Tips:</b> Pastikan format file sesuai dengan template
    </p>
</div>
""", unsafe_allow_html=True)

# --- TAMPILAN MENU 1: REKAP PEMOTONGAN ---
if menu_pilihan == "Rekap Pemotongan":
    # Header section
    col1, col2 = st.columns([3, 1])
    with col1:
        st.title("📝 Rekap Pemotongan")
        st.markdown("Upload file Excel atau CSV untuk memproses data pemotongan")
    with col2:
        st.markdown("<br>", unsafe_allow_html=True)
        with st.expander("ℹ️ Panduan"):
            st.markdown("""
            **Format file yang didukung:**
            - Excel (.xlsx)
            - CSV (.csv)
            
            **Cara menggunakan:**
            1. Pilih file dari komputer
            2. Tunggu proses transformasi
            3. Download hasil Excel
            """)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Upload section
    uploaded_file_rekap = st.file_uploader(
        "📁 Pilih file Excel atau CSV", 
        type=['xlsx', 'csv'], 
        key="rekap_pemotongan",
        help="Upload file data mentah untuk diproses"
    )
    
    if uploaded_file_rekap:
        # Processing section
        with st.spinner('🔄 Memproses file... Mohon tunggu sebentar'):
            st.info(f"✅ File diterima: **{uploaded_file_rekap.name}**")
            result_df_rekap = transform_rekap_pemotongan(uploaded_file_rekap)
        
        if result_df_rekap is not None:
            # Success message
            st.success("🎉 Transformasi data berhasil!")
            
            # Metrics
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("📊 Total Baris", len(result_df_rekap))
            with col2:
                st.metric("📋 Total Kolom", len(result_df_rekap.columns))
            with col3:
                st.metric("📁 Nama File", uploaded_file_rekap.name[:20] + "...")
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            # Preview data dengan styling
            st.subheader("🕶️ Preview Data")
            st.dataframe(
                result_df_rekap, 
                use_container_width=True,
                height=400
            )
            
            # Generate Excel
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                result_df_rekap.to_excel(writer, index=False, sheet_name='Hasil Rekap Pemotongan', 
                                        startrow=1, startcol=0, header=False)
                format_rekap_pemotongan_excel(writer, result_df_rekap)
            excel_data = output.getvalue()
            now = datetime.datetime.now()
            download_filename = now.strftime("%d_%m_%Y-%H_%M") + "-Rekap_Pemotongan.xlsx"
            
            # Download section
            st.markdown("<br>", unsafe_allow_html=True)
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                st.download_button(
                    label="⬇️ Download Hasil sebagai Excel",
                    data=excel_data,
                    file_name=download_filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )

# --- TAMPILAN MENU 2: REKAP KEBUTUHAN MINGGUAN ---
elif menu_pilihan == "Rekap Kebutuhan Mingguan":
    # Header section
    col1, col2 = st.columns([3, 1])
    with col1:
        st.title("📊 Rekap Kebutuhan Mingguan")
        st.markdown("Upload **File Status Penjualan** untuk membuat rekap kebutuhan mingguan")
    with col2:
        st.markdown("<br>", unsafe_allow_html=True)
        with st.expander("ℹ️ Panduan"):
            st.markdown("""
            **File yang dibutuhkan:**
            - File Excel Status Penjualan
            
            **Fitur:**
            - Kelola kategori paket & menu
            - Transformasi otomatis
            - Export ke Excel
            """)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # ===== KATEGORI MANAGEMENT (EXPANDER) =====
    with st.expander("⚙️ Kelola Kategori Paket & Menu", expanded=False):
        st.subheader("Kategori Saat Ini")
        df_kategori = load_kategori()
        
        if len(df_kategori) > 0:
            st.dataframe(df_kategori, use_container_width=True, height=250)
        else:
            st.info("Belum ada kategori.")
        
        # Tab untuk operasi kategori
        kat_tab1, kat_tab2, kat_tab3 = st.tabs(["➕ Tambah", "✏️ Edit", "🗑️ Hapus"])
        
        # TAB 1: Tambah Kategori
        with kat_tab1:
            col1, col2 = st.columns(2)
            with col1:
                nama_barang_baru = st.text_input("Nama Barang / Paket & Menu:", key="nama_barang_tambah_kat")
            with col2:
                kategori_final_baru = st.text_input("Kategori Final:", key="kategori_final_tambah_kat")
            
            if st.button("✅ Tambah", key="btn_tambah_kat"):
                if nama_barang_baru.strip() and kategori_final_baru.strip():
                    df_kat_check = load_kategori()
                    if nama_barang_baru in df_kat_check['Nama Barang'].values:
                        st.warning("⚠️ Nama barang sudah ada.")
                    else:
                        add_kategori(nama_barang_baru, kategori_final_baru)
                        st.rerun()
                else:
                    st.error("❌ Tidak boleh kosong!")
        
        # TAB 2: Edit Kategori
        with kat_tab2:
            df_kat_edit = load_kategori()
            if len(df_kat_edit) > 0:
                idx_edit = st.selectbox("Pilih kategori untuk diedit:", 
                                       range(len(df_kat_edit)), 
                                       format_func=lambda i: f"{i}. {df_kat_edit.iloc[i]['Nama Barang']}",
                                       key="select_edit_kat")
                
                current_nama = df_kat_edit.iloc[idx_edit]['Nama Barang']
                current_kategori = df_kat_edit.iloc[idx_edit]['Kategori Final']
                
                col1, col2 = st.columns(2)
                with col1:
                    nama_barang_edit = st.text_input("Nama Barang:", value=current_nama, key="nama_barang_edit_kat")
                with col2:
                    kategori_final_edit = st.text_input("Kategori Final:", value=current_kategori, key="kategori_final_edit_kat")
                
                if st.button("💾 Simpan", key="btn_edit_kat"):
                    if nama_barang_edit.strip() and kategori_final_edit.strip():
                        update_kategori(idx_edit, nama_barang_edit, kategori_final_edit)
                        st.rerun()
                    else:
                        st.error("❌ Tidak boleh kosong!")
            else:
                st.info("Tidak ada kategori untuk diedit.")
        
        # TAB 3: Hapus Kategori
        with kat_tab3:
            df_kat_delete = load_kategori()
            if len(df_kat_delete) > 0:
                idx_hapus = st.selectbox("Pilih kategori untuk dihapus:", 
                                        range(len(df_kat_delete)), 
                                        format_func=lambda i: f"{i}. {df_kat_delete.iloc[i]['Nama Barang']}",
                                        key="select_delete_kat")
                
                st.warning(f"⚠️ Akan dihapus: **{df_kat_delete.iloc[idx_hapus]['Nama Barang']}**")
                
                if st.button("🗑️ Hapus", key="btn_hapus_kat"):
                    delete_kategori(idx_hapus)
                    st.rerun()
            else:
                st.info("Tidak ada kategori untuk dihapus.")
        
        # Backup & Restore
        st.divider()
        col1, col2 = st.columns(2)
        with col1:
            df_kat_backup = load_kategori()
            if st.button("⬇️ Download kategori.csv"):
                csv_data = df_kat_backup.to_csv(index=False)
                st.download_button(
                    label="📥 Download",
                    data=csv_data,
                    file_name="kategori_backup.csv",
                    mime="text/csv"
                )
        
        with col2:
            uploaded_kategori = st.file_uploader("📤 Upload kategori.csv:", type=['csv'], key="upload_kategori_kat")
            if uploaded_kategori is not None:
                try:
                    df_upload = pd.read_csv(uploaded_kategori)
                    if 'Nama Barang' in df_upload.columns and 'Kategori Final' in df_upload.columns:
                        save_kategori(df_upload)
                        st.rerun()
                    else:
                        st.error("❌ Format file salah!")
                except Exception as e:
                    st.error(f"❌ Error: {e}")
    
    st.divider()
    
    # ===== MAIN PROCESSING =====
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Upload section
    uploaded_file_sales = st.file_uploader(
        "📁 Pilih File Excel Status Penjualan",
        type=['xlsx'],
        key="status_penjualan",
        help="Upload file Excel dengan data status penjualan"
    )
    
    if uploaded_file_sales:
        # Processing section
        with st.spinner('🔄 Memproses data penjualan... Mohon tunggu sebentar'):
            st.info(f"✅ File diterima: **{uploaded_file_sales.name}**")
            result_df_kebutuhan = transform_rekap_kebutuhan(uploaded_file_sales)
        
        if result_df_kebutuhan is not None:
            # Success message
            st.success("🎉 Transformasi data berhasil!")
            
            # Metrics
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("📊 Total Item", len(result_df_kebutuhan))
            with col2:
                st.metric("📋 Total Periode", len(result_df_kebutuhan.columns) - 1)
            with col3:
                st.metric("📁 File", uploaded_file_sales.name[:20] + "...")
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            # Preview data
            st.subheader("🕶️ Preview Rekap Kebutuhan")
            st.dataframe(
                result_df_kebutuhan,
                use_container_width=True,
                height=400
            )
            
            # Generate Excel
            output_kebutuhan = io.BytesIO()
            with pd.ExcelWriter(output_kebutuhan, engine='xlsxwriter') as writer:
                result_df_kebutuhan.to_excel(writer, index=False, sheet_name='Rekap Kebutuhan', startrow=1, header=False)
                workbook = writer.book
                worksheet = writer.sheets['Rekap Kebutuhan']
                header_format = workbook.add_format({
                    'bold': True, 
                    'text_wrap': True, 
                    'valign': 'vcenter', 
                    'align': 'center',
                    'fg_color': '#D7E4BC', 
                    'border': 1
                })
                for col_num, value in enumerate(result_df_kebutuhan.columns.values):
                    worksheet.write(0, col_num, value, header_format)
                # Set row height untuk header agar cukup untuk 2 baris teks
                worksheet.set_row(0, 30)
                worksheet.set_column(0, 0, 40)
                worksheet.set_column(1, len(result_df_kebutuhan.columns)-1, 15)
            excel_data_kebutuhan = output_kebutuhan.getvalue()
            now = datetime.datetime.now()
            download_filename_kebutuhan = now.strftime("%d_%m_%Y-%H_%M") + "-Rekap_Kebutuhan_Mingguan.xlsx"
            
            # Download section
            st.markdown("<br>", unsafe_allow_html=True)
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                st.download_button(
                    label="⬇️ Download Rekap Kebutuhan Mingguan",
                    data=excel_data_kebutuhan,
                    file_name=download_filename_kebutuhan,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )

# --- TAMPILAN MENU 3: LABEL MASAK ---
elif menu_pilihan == "Label Masak":
    # Header section
    col1, col2 = st.columns([3, 1])
    with col1:
        st.title("🏷️ Label Masak")
        st.markdown("Upload file template Excel untuk membuat label masak dalam format Microsoft Word")
    with col2:
        st.markdown("<br>", unsafe_allow_html=True)
        with st.expander("ℹ️ Panduan"):
            st.markdown("""
            **File yang dibutuhkan:**
            - Template Excel dengan data label
            
            **Output:**
            - Dokumen Word (.docx)
            - Format siap cetak
            - Layout otomatis
            """)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Upload section
    uploaded_file_label = st.file_uploader(
        "📁 Pilih File Template Excel",
        type=['xlsx'],
        key="label_masak",
        help="Upload template Excel untuk membuat label masak"
    )
    
    if uploaded_file_label is not None:
        # Processing with better visual feedback
        with st.spinner('🔄 Membuat dokumen Word... Mohon tunggu sebentar'):
            progress_bar = st.progress(0)
            progress_bar.progress(30)
            word_file_buffer = transform_and_create_word_label(uploaded_file_label)
            progress_bar.progress(100)
        
        if word_file_buffer is not None:
            # Success message
            st.success("✅ Dokumen Word berhasil dibuat!")
            
            # Info card
            st.info(f"📄 File template: **{uploaded_file_label.name}** berhasil diproses")
            
            # Generate filename
            now = datetime.datetime.now()
            download_filename_word = now.strftime("%d_%m_%Y-%H_%M") + "-Label_Masak.docx"
            
            # Download section
            st.markdown("<br>", unsafe_allow_html=True)
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                st.download_button(
                    label="⬇️ Download Label Masak sebagai Word",
                    data=word_file_buffer,
                    file_name=download_filename_word,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            # Additional info
            st.markdown("<br>", unsafe_allow_html=True)
            with st.expander("📝 Informasi Dokumen"):
                st.markdown(f"""
                - **Nama File**: {download_filename_word}
                - **Format**: Microsoft Word (.docx)
                - **Dibuat**: {now.strftime("%d %B %Y, %H:%M")}
                - **Status**: Siap untuk dicetak
                """)